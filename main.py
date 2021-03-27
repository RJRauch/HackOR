#################################################################################################
#                                                                                               #
#       Batch Editor (insert fun name here) - a Python Program Written for HackOR 2021          #
#       Written by:                                                                             #
#           Sophia L                                                                            #
#           Ryan R.                                                                             #
#           Noah S.                                                                             #
#                                                                                               #
#################################################################################################

# necessary libraries


import os                   # for file navigation
import docx                 # to read docX files
from pathlib import Path    # for creating paths


def prompt_user() -> tuple:
    """prompt_user - prompts the user to enter an absolute filepath

    @param: None

    @return: A tuple containing [0] a string representing the folder path, and [1] a dictionary containing the actions to be performed
    @rtype : Tuple
    """

    actions = {}

    # get the path from the user

    print("Hello, welcome to Batch Formatter v0.1 Beta. Please enter the absolute path of the folder where the text files you would like to reformat are located.")
    absolute_path = input()

    # determine if the user would like to
    header_text = get_text("header")
    if header_text != None:
        actions["1"] = header_text

    # determine if the user wants to add a footer
    footer_text = get_text("footer")
    if footer_text != None:
        actions["2"] = footer_text

    # determine if the user wants to add page numbers
    add_page_numbers = get_binary_answer("add page numbers")
    if add_page_numbers == True:
        actions["3"] = True

    return (absolute_path, actions)


def get_text(sought_text_type: str) -> str:
    """get_text - prompts the user to enter the text they want for the header - returns none if they are not interested

    @param sought_text_type:    A description of the purpose you are seeking text from the user for
    @param type:                String

    @return: The user's string if one was entered, else None
    @rtype : str
    """

    # prompt the user to see if they are interested in adding a header
    print("Would you like to enter a " + sought_text_type + "? [Y/N]")
    answer = input()

    # call a validator to see if they need to re-enter
    valid_answer = validate_user_answer(answer)

    if not valid_answer:
        get_text(sought_text_type)

    if valid_answer == 'y' or valid_answer == 'yes':
        print("What would you like to add as your " + sought_text_type + "?")
        text = input()
        return text

    return None


def get_binary_answer(answer_topic: str) -> bool:
    """get_binary_answer - prompts the user to enter whether or not they are interested in a given topic and returns an answer

    @param answer_topic : The topic that you are seeking an answer from the user on
    @param type         : String

    @return: A bool indicating if the user is interested in the topic
    @rtype : Bool
    """

    # prompt the user on the subject
    print("Would you like to " + answer_topic + "? [Y/N]")
    answer = input()

    # validate their input
    valid_answer = validate_user_answer(answer)
    if valid_answer == None:
        get_binary_answer(answer_topic)

    # once valid input has been entered return a bool based on their preference
    if valid_answer.lower() == 'y' or valid_answer.lower() == 'yes':
        return True

    return False


def validate_user_answer(answer: str) -> str:
    """answer_is_valid - validates user input gauranteeing it is either yes, no, or some parseable variation on one of them.

    @param answer: The user's input
    @type  answer: String

    @return: Returns the answer if it was valid, else None
    @rtype : String
    """

    # if the answer is yes or no return true
    if answer.lower() == 'y' or answer.lower() == 'n':
        return answer.lower()

    if answer.lower() == 'yes' or answer.lower() == 'no':
        return answer.lower()

    else:
        print("You must enter [Y]es or [N]o")
        return None


def traverse_directory(action_tuple: tuple) -> None:
    """traverse_directory - takes an input tuple containing a directory to traverse and actions to perform - reads all files from the directory and passes them to perform modification with the actions specified

    @param action_tuple: A tuple containing the directory to traverse and the actions to execute
    @type  action_tuple: tuple

    @return: None
    """

    directory = action_tuple[0]
    action_dict = action_tuple[1]
    if len(action_dict) == 0:
        print("No Actions Selected")
        return

    actions = list(action_tuple[1].keys())

    for folderName, subfolders, filenames in os.walk(directory):

        os.chdir(Path(folderName))
        # create an absolute file path for each file we need to edit
        for filename in filenames:
            # for each action contained in the actions array, alert the user of what you will do then perform that action
            if "1" in actions:
                print("adding header to: " + filename)
                add_header(filename, action_dict['1'])
            if "2" in actions:
                print("adding footer to: " + filename)
                add_footer(filename, action_dict['2'])
            if "3" in actions:
                print("adding page numbers to: " + filename)
                add_page_numbers(filename)


def add_header(file: str, header_text: str) -> None:
    """Mutates doc, docx, and txt files by adding a custom header.

    @param file: The file path of the file to add the header to
    @type  file: String

    @param header_text: The header text to be added to each file
    @type  header_text: String

    @return: None
    """

    # Checks if doc or docx file, adds header to it
    if ".doc" in file or ".docx" in file:
        doc = docx.Document(file)
        doc.add_heading(header_text, 0)

    # Checks if txt file, adds header to it
    elif ".txt" in file:
        doc = open(file, "w+")  # w+ puts cursor at beginning of file
        doc.write(header_text + "\n")
        doc.close()

    else:
        print("File must be .doc, .docx, or .txt format.")

    return None


def add_footer(file: str, footer_text: str) -> None:
    """
    This function will add a footer to .docx/doc/txt file

    @param file: the file to add the footer to
    @Type file: string
    @Param footer_Text: the footer text to be added to each file
    @ type footer_text: string'''

    @return: None
    """

    # Check for  doc/docx and add footer to it
    if '.doc' in file or '.docx' in file:
        # Create footer for document object
        doc = docx.Document(file)

        # footer section
        footer = doc.sections[0].footer
        footer.paragraphs[0].text = footer_text
        # close file
        doc.save(file)

        # txt document This will technically create new lines from paragraph text and just add to the end of the document
        # It will not create consistent page footers as there are not pages in a text document

    elif '.txt' in file:
        f_hand = open(file, 'a+')
        f_hand.write(footer_text)
        f_hand.close()

    else:
        print(
            'Error: You have to enter a text or doc/docx file or this will never EVER work')


def add_page_numbers(file: str) -> None:
    """add_page_numbers- adds page numbers to each page of a file

    @param file: the file to add the page numbers to
    @type  file: String
    """
    pass


def main():
    actions = prompt_user()
    traverse_directory(actions)
    print("Files Modified")


if __name__ == "__main__":
    main()
